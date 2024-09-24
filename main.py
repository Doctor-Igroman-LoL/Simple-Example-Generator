from datetime import date

import docx
from docx.shared import Pt
from docx.shared import Cm
from docx.shared import RGBColor
import random
from openpyxl import load_workbook

doc = docx.Document()
sections = doc.sections
for section in sections:
    section.top_margin = Cm(0.5)
    section.bottom_margin = Cm(0.5)
    section.left_margin = Cm(1)
    section.right_margin = Cm(1)

wb = load_workbook('list_students.xlsx')

array_students = []
array_all_answer = []
data_answer = []
count_example = 4

level_addition = 1          # Сложение
level_subtraction = 2       # Вычитание
level_multiplication = 2    # Умножение
level_division = 1         # Деление


def calculate_the_list_of_students(number_students=1):
    sheet = wb['Лист1']
    for i in range(number_students):
        array_students.append(sheet.cell(row=i+1, column=1).value)



def write_docx(count_example=4):
    calculate_the_list_of_students(11)
    n = 1
    for student in array_students:

        name_student = student
        data_answer.append(name_student)
        _date = "\t\t\t\t\t\t\t\t\t\t\t\t[" + str(date.today()) + "]\n"
        stroka = doc.add_paragraph(_date + 'Здравствуй {}. Реши математические примеры:'.format(name_student))
        stroka.paragraph_format.space_after = Pt(1)
        #stroka = doc.add_paragraph('### Умножение')  # ### Сложение
        #stroka.paragraph_format.space_after = Pt(1)
        #generation_example('*', count_example, level_addition)
        #stroka = doc.add_paragraph('### Умножение')  # ### Вычитание
        #stroka.paragraph_format.space_after = Pt(1)
        #generation_example('*', count_example, level_subtraction)
        stroka = doc.add_paragraph('### Деление')
        stroka.paragraph_format.space_after = Pt(1)
        generation_example('/', count_example, level_multiplication)
        #stroka = doc.add_paragraph('### Деление')
        #stroka.paragraph_format.space_after = Pt(1)
       # generation_example('/', count_example, level_division)

        stroka = doc.add_paragraph('')
        stroka.paragraph_format.space_after = Pt(5)
        n += 1

        if n % 6 == 0 and n != 0:
            print(n)
            doc.add_page_break()

    array_all_answer.append(data_answer)

    stroka = doc.add_paragraph("")

    for _l in array_all_answer[0]:
        stroka.add_run(str(_l) + " ")
        stroka.paragraph_format.space_after = Pt(5)

    # Сохраняем записи в файле
    doc.save('examples.docx')


# Функция генерация математических примеров
def generate_simple_expression(operator, _level):
    # Операции ['+', '-', '*', '/']
    answer = 0
    # Уровни примеров
    ten = 1
    if _level != 1:
        for _ in range(_level-1):
            ten *= 10

    range_numbers = 10 * ten

    # Генерация случайных чисел
    num1 = random.randint(2, range_numbers)
    num2 = random.randint(1, range_numbers)

    # Обработка вычитание для избежания отрицательных примеров
    if operator == '+':
        answer = num1 + num2

    # Обработка вычитание для избежания отрицательных примеров
    if operator == '-':
        num1 = num1 + num2
        answer = num1 - num2

    # Обработка деления для избежания деления на ноль
    if operator == '/':
        num1 = num1 * num2  # Убедимся, что делитель не ноль и результат будет целым
        answer = num1 / num2
        operator = ':'

    # Если умножение, заменяем знак на точку
    if operator == '*':
        answer = num1 * num2
        operator = '·'

    # Формирование примера
    expression = f"{num1} {operator} {num2} = "

    return [expression, answer]


# Генерация и вывод нескольких примеров
def generation_example(operator, count, level=1):
    stroka = doc.add_paragraph()

    _text = " (" + operator + ") "
    data_answer.append(_text)

    for i in range(count):
        result = generate_simple_expression(operator, level)
        example = str(i + 1) + ") " + result[0]
        stroka.add_run(example)
        answer = result[1]
        data_answer.append(str(int(answer)) + ", ")

        if i == range(count):
            stroka.add_run("\n")
            stroka.paragraph_format.space_after = Pt(6)

        if (i + 1) % 4 == 0:
            stroka.add_run("\n")

        else:
            stroka.add_run("\t\t")


write_docx(count_example)
print("Выполнено хозяйн ٩(◕‿◕｡)۶")